import os
# import cv2
import math
import numpy as np
from docx import Document
from cert_imgs_db import CertImagesDB
from docx.shared import Cm
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor#设置字体

from PIL import Image, ImageDraw, ImageFont

cert_imgs_path = ''


def cert_img_to_doc(records: list, out_dir: str, watermark_text: str):
    if not os.path.exists(out_dir) or not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    document = Document()
    if list is None and len(records) > 0:
        return document;
    # document.add_heading('资质证书列表')
    for record in records:
        name = record['name']
        imgs = record['images']
        title = document.add_heading('', level=3)
        title_run = title.add_run(name)
        title_run.font.name = '宋体'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_run.font.size = Pt(10)
        title_run.font.bold = True

        try:
            for index, img in enumerate(imgs):
                img_path = os.path.join(os.getcwd(), 'data', 'cert-image', img[len('/cert-imgs/'):])
                if not os.path.exists(img_path):
                    document.add_paragraph('图片不存在:' + img)
                    continue
                img_out_path = os.path.join(out_dir, 'img-{}-{}.png'.format(name, index))
                img = img_process(img_path, save_path=img_out_path, watermark_text=watermark_text)
                document.add_picture(img_out_path, width=Cm(15))
        except Exception as e:
            document.add_paragraph('插入失败:' + name)
            print(e)
        document.add_page_break()

    return document


# 处理文档中的图片（翻转，缩放，水印）
def img_process(img_path, save_path=None, watermark_text: str = None):
    max_w = 1200
    max_h = 1500
    if not os.path.exists(img_path):
        return None

    # 翻转
    img = Image.open(img_path)
    if img is None:
        return None
    w, h = img.size[:2]
    if h < w:
        # img = cv2.rotate(img, rotateCode=cv2.ROTATE_90_COUNTERCLOCKWISE)
        # img = img.rotate(angle=90)
        img = img.transpose(Image.ROTATE_90)
        w, h = img.size[:2]
    if h / max_h > w / max_w:
        new_h = max_h
        new_w = int((w / h) * new_h)
    else:
        new_w = max_w
        new_h = int((h / w) * new_w)

    # 缩放
    img = img.resize((new_w, new_h))

    # 加水印
    if watermark_text is not None and len(watermark_text.strip()) > 0:
        img = img.convert('RGBA')
        logo = gen_text_watermark(text=watermark_text, text_size=30, text_color=(100, 100, 100))
        # img.paste(logo, (0, 0), logo.split()[3])
        logo_full = Image.new(mode='RGBA', size=img.size)

        # 水印平铺
        img_w, img_h = img.size
        logo_w, logo_h = logo.size
        x_cnt = math.ceil(img_w / logo_w)
        y_cnt = math.ceil(img_h / logo_h)
        for x in range(x_cnt):
            for y in range(y_cnt):
                logo_full.paste(logo, (x * logo_w, y * logo_h))
        # 叠加水印
        img = Image.alpha_composite(img, logo_full)

    if save_path is not None:
        img.save(save_path)
    return img


def gen_text_watermark(text='logo', text_size=10, text_color=(100, 100, 100), bg_color=(255, 255, 255, 50)):
    char_cnt = len(text)
    width, height = char_cnt * text_size, text_size
    img = Image.new(mode='RGBA', size=(width, height), color=bg_color)
    draw = ImageDraw.Draw(img)
    # 字体的格式
    font_style = ImageFont.truetype(
        "data/tcc/simsun.ttc", text_size, encoding="utf-8")
    # 绘制文本
    draw.text(xy=(0, 0), text=text, fill=text_color, font=font_style)
    img = img.rotate(angle=45, expand=1, fillcolor=bg_color)
    # img.show()
    # print(img.size)
    return img


def save_doc(document: Document, doc_path: str):
    document.save(doc_path)
    return doc_path


if __name__ == "__main__":
    img_path = 'data/cert-image/b.jpeg'
    # img = cv2.imread(img_path)
    img = img_process(img_path, watermark_text='这是一个水印哦')
    img.show()
    gen_text_watermark(text="你 好");
