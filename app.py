from docx import Document
from cert_imgs_db import CertImagesDB
from docx.shared import Cm
from docx.enum.text import WD_BREAK
import gradio as gr

document = Document()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

document.save('text.docx')

if __name__ == '__main__':
    document = Document()
    document.add_heading('我是一级标题')
    document.add_heading('我是一级标题')
    document.add_heading('我是一级标题')
    document.add_heading('我是二级标题', level=2)
    document.add_page_break()
    document.add_heading('我是二级标题', level=2)
    document.add_picture('./data/cert-image/a.png', width=Cm(10))
    document.add_page_break()
    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
    paragraph.runs[-1].add_break(WD_BREAK.PAGE)

    document.save('text.docx')

    db = CertImagesDB()
    # ret = db.add({
    #     'name': 'ssss',
    #     'images': ['aa.png', 'bb.png']
    # })
    ret = db.query('资质')
    print(ret)
    gr.blocks
